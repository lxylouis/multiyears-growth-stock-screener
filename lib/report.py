"""
multiyears-growth-stock-screener/lib/report.py
共享DOCX报告生成器 —— 统一所有指数的报告模板

所有"方案F"已替换为实际规则描述，报告中不再使用代号。
"""
import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class ReportBuilder:
    """统一的DOCX报告构建器"""

    def __init__(self, index_name: str, subtitle: str = ''):
        self.doc = Document()
        self.index_name = index_name
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.3
        for lv in range(1, 5):
            hs = self.doc.styles[f'Heading {lv}']
            hs.font.name = 'Microsoft YaHei'
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ── 内容元素 ──

    def cover(self, subtitle='', stats=''):
        """封面页"""
        for _ in range(4): self.doc.add_paragraph()
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{self.index_name} · 增长分析报告')
        r.bold = True; r.font.size = Pt(24)
        r.font.name = 'Microsoft YaHei'; r.font.color.rgb = RGBColor(30, 60, 120)

        if subtitle:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle)
            r.font.size = Pt(14); r.font.name = 'Microsoft YaHei'
            r.font.color.rgb = RGBColor(100, 100, 100)

        if stats:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stats)
            r.font.size = Pt(10); r.font.name = 'Microsoft YaHei'
            r.font.color.rgb = RGBColor(130, 130, 130)

        self.doc.add_page_break()

    def heading(self, text, level=1):
        h = self.doc.add_heading(text, level=level)
        for run in h.runs: run.font.name = 'Microsoft YaHei'
        return h

    def para(self, text, bold=False, size=10.5, align=None, color=None):
        p = self.doc.add_paragraph()
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(text)
        rn.font.name = 'Microsoft YaHei'; rn.font.size = Pt(size); rn.bold = bold
        if color: rn.font.color.rgb = RGBColor(*color)
        return p

    def table(self, df, caption=None, font_size=8.5):
        """插入格式化表格"""
        if caption:
            p = self.doc.add_paragraph()
            r = p.add_run(caption)
            r.bold = True; r.font.size = Pt(10); r.font.name = 'Microsoft YaHei'
        rows, cols = df.shape
        table = self.doc.add_table(rows=rows+1, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, col in enumerate(df.columns):
            cell = table.cell(0, j); cell.text = str(col)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(font_size); r.font.name = 'Microsoft YaHei'
        for i in range(rows):
            for j in range(cols):
                cell = table.cell(i+1, j)
                v = df.iloc[i, j]
                text = str(v) if not (isinstance(v, float) and np.isnan(v)) else ''
                cell.text = text
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(font_size); r.font.name = 'Microsoft YaHei'
        return table

    def image(self, path, width=6.3):
        if os.path.exists(path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(width))

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path: str):
        """保存DOCX"""
        self.doc.save(path)
        return path


# ── 标准报告章节工厂 ──

def build_price_report(index_name: str, total: int, passed: int, t20, pp, fp_df,
                       mult_cols, period_labels, chart_paths: dict, out_path: str,
                       cond_a_min: float = 1.3, cond_b_threshold: float = 1.5,
                       window_years: int = 3, num_windows: int = 3):
    """构建统一的股价/营收分析报告"""
    mode = '营收' if '营收' in out_path else '股价'
    rpt = ReportBuilder(index_name, subtitle=f'{mode}增长 · {window_years}年滚动倍率 · 3月末版')
    rpt.cover(subtitle=f'{mode}增长 · {window_years}年滚动倍率 · 3月末版',
              stats=f'保留条件：全期≥{cond_a_min} ∪ 最新期>{cond_b_threshold} | '
                    f'{window_years}年×{num_windows}期 线性递减排序 | 含{total}只成分股')

    # 一、核心结论
    rpt.heading('一、核心结论', 1)
    rpt.para(f'筛选通过 {passed}/{total} 只（{passed/total*100:.0f}%）')
    rpt.para(f'保留条件：过去每期增长倍率均≥{cond_a_min}（每{window_years}年至少增长{((cond_a_min**(1/window_years))-1)*100:.1f}%），'
             f'或最新一期倍率>{cond_b_threshold}（近{window_years}年增长{(cond_b_threshold**(1/window_years))-1:.1%}）')
    rpt.para(f'时间窗口：{window_years}年×{num_windows}期，线性递减权重（越近越高）')
    rpt.para('▎ TOP5 股票', bold=True, size=12, color=(20, 80, 160))
    for i, (_, r) in enumerate(t20.head(5).iterrows(), 1):
        rpt.para(f'  {i}. {r["名称"]} — {r["分数"]:.2f}pt — 年化{r.get("年均化", 0):.1f}%')

    # 二、TOP20表
    rpt.page_break()
    rpt.heading('二、TOP20 完整排名', 1)
    disp = t20[['代码', '名称', '分数', '年均化', '波动', '行业'] + mult_cols].copy()
    disp.columns = ['代码', '名称', '分数', '年化(%)', '波动', '行业'] + period_labels
    disp['年化(%)'] = disp['年化(%)'].apply(lambda x: f"{x:.1f}")
    rpt.table(disp, f'TOP20排名（通过筛选共{passed}只）')

    # 三-七：图表
    for chap, chart_key, chart_title in [
        ('三', 'trend', 'TOP20增长趋势图'),
        ('四', 'heatmap', '热力图'),
        ('五', 'industry', '行业分析'),
        ('六', 'boxplot', '各周期分布（箱线图）'),
        ('七', 'sustainability', '持续性 vs 分数'),
    ]:
        rpt.page_break()
        rpt.heading(f'{chap}、{chart_title}', 1)
        cp = chart_paths.get(chart_key)
        if cp and os.path.exists(cp):
            rpt.image(cp)

    # 八、全量排名（限100行，太长会导致DOCX卡死）
    rpt.page_break()
    rpt.heading('八、全量排名（TOP100）', 1)
    all_disp = fp_df.sort_values('分数', ascending=False).head(100)[['代码', '名称', '分数', '行业'] + mult_cols].copy()
    all_disp['年均化'] = all_disp[mult_cols].mean(axis=1).apply(
        lambda x: f"{(x**(1/len(mult_cols))-1)*100:.1f}%")
    all_disp = all_disp[['代码', '名称', '分数', '年均化', '行业'] + mult_cols]
    all_disp.columns = ['代码', '名称', '分数', '年化(%)', '行业'] + period_labels
    all_disp['年化(%)'] = all_disp['年化(%)'].apply(str)
    rpt.table(all_disp, '全量排名（按分数降序）')

    return rpt.save(out_path)


def build_dual_report(index_name: str, price_pp, rev_pp, price_df, rev_df, merged,
                      price_mc, rev_mc, period_labels, chart_paths: dict, out_path: str,
                       cond_a_min: float = 1.3, cond_b_threshold: float = 1.5,
                       window_years: int = 3, num_windows: int = 3):
    """构建统一的股价+营收双维度分析报告"""
    elite = merged[merged['双通过']].sort_values('综合分数', ascending=False).reset_index(drop=True)
    only_price = merged[merged['仅股价']].sort_values('股价分数', ascending=False)
    only_rev = merged[merged['仅营收']].sort_values('营收分数', ascending=False)

    stats = (f'股价保留 {len(price_pp)}/{len(price_df)}只 | '
             f'营收保留 {len(rev_pp)}/{len(rev_df)}只 | '
             f'双维度 {len(elite)}只')
    rpt = ReportBuilder(index_name, subtitle='股价 + 营收 双维度对比')
    rpt.cover(subtitle='股价 + 营收 双维度对比', stats=stats)

    # 一、方法论
    rpt.heading('一、方法论概览', 1)
    rpt.para(f'数据源：股价=每年3月末收盘价，营收=年报营业总收入')
    rpt.para(f'时间窗口：{window_years}年×{num_windows}期滚动')
    rpt.para(f'筛选规则：保留条件 = 全期≥{cond_a_min} ∪ 最新期>{cond_b_threshold}')
    rpt.para(f'  · 条件A：过去每一期增长倍率均不低于{cond_a_min}（每{window_years}年至少增长{((cond_a_min**(1/window_years))-1)*100:.1f}%）')
    rpt.para(f'  · 条件B：最新一期增长倍率超过{cond_b_threshold}（近{window_years}年增长{(cond_b_threshold**(1/window_years))-1:.1%}）')
    rpt.para(f'评分：{num_windows}期线性递减权重，越近的周期权重越高（1:2:...:{num_windows}）')
    rpt.para('')

    # 对比表
    ct = rpt.doc.add_table(rows=6, cols=3)
    ct.style = 'Table Grid'; ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['对比项', '股价', '营收']):
        cell = ct.cell(0, j); cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9.5)
    for i, (item, v1, v2) in enumerate([
        ('有效样本', f'{len(price_df)}只', f'{len(rev_df)}只'),
        ('保留（通过筛选）', f'{len(price_pp)}只 ({len(price_pp)/len(price_df)*100:.0f}%)',
                            f'{len(rev_pp)}只 ({len(rev_pp)/len(rev_df)*100:.0f}%)'),
        ('分数均值', f'{price_pp["分数"].mean():.2f}pt', f'{rev_pp["分数"].mean():.2f}pt'),
        ('最高分数', f'{price_pp["分数"].max():.2f}pt', f'{rev_pp["分数"].max():.2f}pt'),
        ('特征', '股价驱动型', '营收驱动型'),
    ], 1):
        for j, v in enumerate([item, v1, v2]):
            cell = ct.cell(i, j); cell.text = v
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)

    # 二、双维度精选
    rpt.page_break()
    rpt.heading('二、双维度精选', 1)
    rpt.para(f'共 {len(elite)} 只股票同时通过股价和营收筛选保留')
    if len(elite) > 0:
        ed = elite[['代码', '名称', '行业', '股价分数', '营收分数', '综合分数']].round(2).copy()
        rpt.table(ed, f'双维度精选（{len(elite)}只）')

    # 三、股价TOP20
    rpt.page_break()
    rpt.heading('三、股价增长 TOP20', 1)
    pt20 = price_pp.head(20)
    pd_disp = pt20[['代码', '名称', '分数', '行业'] + price_mc].copy()
    pd_disp['年均化'] = pt20['年均化'].apply(lambda x: f"{x:.1f}")
    pd_disp = pd_disp[['代码', '名称', '分数', '年均化', '行业'] + price_mc]
    pd_disp.columns = ['代码', '名称', '分数', '年化(%)', '行业'] + period_labels
    rpt.table(pd_disp, f'股价TOP20（通过筛选共{len(price_pp)}只）')

    # 四、营收TOP20
    rpt.page_break()
    rpt.heading('四、营收增长 TOP20', 1)
    rt20 = rev_pp.head(20)
    rd_disp = rt20[['代码', '名称', '分数', '行业'] + rev_mc].copy()
    rd_disp['年均化'] = rt20['年均化'].apply(lambda x: f"{x:.1f}")
    rd_disp = rd_disp[['代码', '名称', '分数', '年均化', '行业'] + rev_mc]
    rd_disp.columns = ['代码', '名称', '分数', '年化(%)', '行业'] + period_labels
    rpt.table(rd_disp, f'营收TOP20（通过筛选共{len(rev_pp)}只）')

    # 五、对比分析
    rpt.page_break()
    rpt.heading('五、双维度对比分析', 1)

    for cp_key, title in [('dual_bar', '股价 vs 营收 对比图'), ('dual_scatter', '双维度散点图')]:
        cp = chart_paths.get(cp_key)
        if cp and os.path.exists(cp):
            rpt.heading(title, 2)
            rpt.image(cp)

    # 仅股价 / 仅营收
    rpt.para('')
    rpt.para('5.3 仅股价强劲', bold=True, size=12)
    if len(only_price) > 0:
        rpt.table(only_price[['代码', '名称', '股价分数', '营收分数']].round(2),
                  f'仅股价通过（{len(only_price)}只）')
    else:
        rpt.para('  无（所有股价通过股票均同时通过营收筛选）')

    rpt.para('')
    rpt.para('5.4 仅营收强劲', bold=True, size=12)
    if len(only_rev) > 0:
        rpt.table(only_rev[['代码', '名称', '营收分数', '股价分数']].round(2),
                  f'仅营收通过（{len(only_rev)}只）')

    # 六、精选趋势
    if chart_paths.get('dual_trend') and os.path.exists(chart_paths['dual_trend']):
        rpt.page_break()
        rpt.heading('六、双维度精选增长趋势', 1)
        rpt.image(chart_paths['dual_trend'])

    return rpt.save(out_path)
